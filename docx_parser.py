"""
Use this script to parse .docx files downloaded from CONFLUENCE.
Convert downloaded .doc files to .docx with `Save As` in MS Word.
"""

import pandas as pd
from docx import Document


DOCUMENT_PATH = r"../script_io_files/My+home+new.docx"
OUTPUT_EXCEL_PATH = r"../script_io_files/docx_parser_output.xlsx"

output_dict = {"Event name": [], "Description": []}

doc = Document(DOCUMENT_PATH)

outer_table = doc.tables[0]  # selecting the outermost table containing other tables
for outer_row in outer_table.rows:
    event_cell = outer_row.cells[1]  # refers to the cell of the Events column
    if event_cell.tables:
        inner_event_table = event_cell.tables[0]

        for inner_event_row in inner_event_table.rows:

            if len(inner_event_row.cells) >= 2:
                first_cell = inner_event_row.cells[0].text.strip()
                second_cell = inner_event_row.cells[1].text.strip()

                if first_cell.lower() == "event name":
                    event_names = second_cell.split("\n")  # some events have multiple event names separated by a newline
                    output_dict["Event name"].extend(event_names)

                if first_cell.lower() == "description":
                    description_names = second_cell.split("\n")  # not sure if this is necessary for docs
                    output_dict["Description"].extend(description_names)


df = pd.DataFrame(output_dict)
df.index = range(1, len(df)+1)
df.to_excel(OUTPUT_EXCEL_PATH, index=True)